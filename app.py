import streamlit as st 
import requests
from docx import Document
from io import BytesIO, StringIO
import pandas as pd

st.set_page_config(page_title="Asistente IA", layout="wide")

# Configuración de LM Studio
LM_STUDIO_URL = "http://192.168.18.35:1234/v1/chat/completions"
MODEL_NAME = "phi-2"  # Cambia si usas otro modelo

# -------- FUNCIONES --------

def generate_article(topic):
    try:
        headers = {"Content-Type": "application/json"}
        data = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": "Eres un experto en redacción de artículos SEO. Siempre responde en español."},
                {"role": "user", "content": f"Escribe un artículo optimizado para SEO sobre: {topic}"}
            ],
            "temperature": 0.7,
            "max_tokens": 1000
        }
        response = requests.post(LM_STUDIO_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        st.error(f"Error al generar el artículo: {str(e)}")
        return None

def create_word_doc(text, title="Documento Generado"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_code(language, description):
    try:
        headers = {"Content-Type": "application/json"}
        data = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": f"Eres un programador experto. Siempre responde en español. Genera código en {language}."},
                {"role": "user", "content": f"Genera un código en {language} que: {description}"}
            ],
            "temperature": 0.5,
            "max_tokens": 800
        }
        response = requests.post(LM_STUDIO_URL, headers=headers, json=data)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        st.error(f"Error al generar el código: {str(e)}")
        return None

def generate_table(description):
    try:
        headers = {"Content-Type": "application/json"}
        data = {
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": "Eres un experto en análisis de datos. Siempre responde en español. Devuelve solo una tabla en formato CSV con encabezado."},
                {"role": "user", "content": f"Genera una tabla con datos que represente: {description}"}
            ],
            "temperature": 0.5,
            "max_tokens": 1000
        }
        response = requests.post(LM_STUDIO_URL, headers=headers, json=data)
        response.raise_for_status()
        table_csv = response.json()["choices"][0]["message"]["content"].strip()

        # Si devuelve en formato Markdown, limpiarlo
        if table_csv.startswith("|"):
            lines = table_csv.splitlines()
            clean_lines = [line for line in lines if not line.strip().startswith("|---")]
            table_csv = "\n".join([line.strip().strip('|').replace(" | ", ",") for line in clean_lines])

        return table_csv
    except Exception as e:
        st.error(f"Error al generar la tabla: {str(e)}")
        return None

# -------- INTERFAZ --------

st.title('Generador de contenidos IA')

section = st.sidebar.selectbox(
    'Selecciona una sección:',
    ['Generacion de articulos', 'Generacion de codigo', 'Generacion de tablas de Datos']
)

if section == 'Generacion de articulos':
    st.header('Generación de Artículos')
    topic = st.text_input('Ingresa un tema para el artículo:')
    if st.button('Generar Artículo'):
        if topic:
            with st.spinner('Generando artículo...'):
                article = generate_article(topic)
                if article:
                    st.success('Artículo generado exitosamente')
                    st.markdown("### Vista previa del artículo")
                    st.markdown(article)
                    buffer = create_word_doc(article, title="Artículo Generado")
                    st.download_button(
                        label="Descargar como Word",
                        data=buffer,
                        file_name="articulo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
        else: 
            st.warning('Por favor ingresa un tema para el artículo')

elif section == 'Generacion de codigo':
    st.header('Generación de Código')
    language = st.selectbox('Selecciona el lenguaje de programación:', ['Python', 'SQL', 'HTML', 'JavaScript', 'Java', 'C#'])
    description = st.text_area('Describe qué código deseas generar:')
    if st.button('Generar Código'):
        if description:
            with st.spinner('Generando código...'):
                code = generate_code(language, description)
                if code:
                    st.success('Código generado exitosamente')
                    st.markdown("### Vista previa del código")
                    st.code(code, language=language.lower())
                    st.download_button(
                        label="Descargar como archivo .txt",
                        data=BytesIO(code.encode()),
                        file_name=f"codigo_{language.lower()}.txt",
                        mime="text/plain"
                    )
        else:
            st.warning('Por favor describe el código que deseas generar')

elif section == 'Generacion de tablas de Datos':
    st.header('Generación de Tablas de Datos')
    description = st.text_area('Describe la tabla que deseas generar (ej: Productos con precios y stock):')
    if st.button('Generar Tabla'):
        if description:
            with st.spinner('Generando tabla...'):
                table_csv = generate_table(description)
                if table_csv:
                    try:
                        df = pd.read_csv(StringIO(table_csv))
                        st.success("Tabla generada exitosamente")
                        st.markdown("### Vista previa de la tabla")
                        st.dataframe(df)
                        st.download_button(
                            label="Descargar como CSV",
                            data=table_csv.encode(),
                            file_name="tabla_generada.csv",
                            mime="text/csv"
                        )
                    except Exception:
                        st.error("La respuesta no tiene formato válido de tabla CSV.")
                        st.text(table_csv)
        else:
            st.warning('Por favor describe la tabla que deseas generar')
